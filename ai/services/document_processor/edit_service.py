"""
EditService — document read, LLM-based edit proposal generation, and safe file write.

Phase 1 supports .txt and .docx files only.
Every write is preceded by a timestamped .bak copy; no file is touched until
the user explicitly confirms the proposal via the /api/documents/edit/apply endpoint.
"""

import json
import logging
import os
import re
import shutil
import uuid
from dataclasses import dataclass, field
from datetime import datetime, timedelta
from html.parser import HTMLParser
from pathlib import Path
from typing import Any, Dict, List, Optional

logger = logging.getLogger(__name__)

MAX_READ_CHARS = 12_000   # truncate very large files before LLM call
PROPOSAL_TTL_SECONDS = 1800  # proposals expire after 30 minutes
SUPPORTED_EDIT_EXTENSIONS = {".txt", ".docx"}


# ---------------------------------------------------------------------------
# Data structures
# ---------------------------------------------------------------------------

@dataclass
class EditChange:
    find: str     # exact verbatim text to locate in document
    replace: str  # replacement text


@dataclass
class EditProposal:
    proposal_id: str
    document_name: str
    file_path: str
    file_type: str          # "txt" or "docx"
    changes: List[EditChange]
    summary: str
    expires_at: datetime = field(
        default_factory=lambda: datetime.utcnow() + timedelta(seconds=PROPOSAL_TTL_SECONDS)
    )

    def to_client_dict(self) -> Dict[str, Any]:
        """Serialise for the frontend — no internal state included."""
        return {
            "proposal_id": self.proposal_id,
            "document_name": self.document_name,
            "file_type": self.file_type,
            "changes": [{"find": c.find, "replace": c.replace} for c in self.changes],
            "summary": self.summary,
        }


# ---------------------------------------------------------------------------
# Service
# ---------------------------------------------------------------------------

class EditService:
    """
    Handles the full edit lifecycle:
      1. read_for_edit   — extract plain text from file for the LLM
      2. generate_proposal — LLM produces structured find/replace diff
      3. apply_proposal  — validate, back up, write, return result

    `current_directory` must be set by the orchestrator whenever a new
    directory is loaded; it is used as a scope guard before every write.
    """

    def __init__(self) -> None:
        self.current_directory: Optional[str] = None
        self._pending: Dict[str, EditProposal] = {}

    # ------------------------------------------------------------------
    # Proposal store
    # ------------------------------------------------------------------

    def store_proposal(self, proposal: EditProposal) -> None:
        self._purge_expired()
        self._pending[proposal.proposal_id] = proposal

    def get_proposal(self, proposal_id: str) -> Optional[EditProposal]:
        self._purge_expired()
        return self._pending.get(proposal_id)

    def remove_proposal(self, proposal_id: str) -> None:
        self._pending.pop(proposal_id, None)

    def _purge_expired(self) -> None:
        now = datetime.utcnow()
        expired = [pid for pid, p in self._pending.items() if p.expires_at < now]
        for pid in expired:
            del self._pending[pid]

    # ------------------------------------------------------------------
    # Read
    # ------------------------------------------------------------------

    def can_edit(self, file_path: str) -> bool:
        return Path(file_path).suffix.lower() in SUPPORTED_EDIT_EXTENSIONS

    def read_for_edit(self, file_path: str) -> Optional[str]:
        """Return a plain-text snapshot of the file for LLM consumption."""
        path = Path(file_path)
        if not path.exists() or not path.is_file():
            logger.warning("read_for_edit: not found: %s", file_path)
            return None
        suffix = path.suffix.lower()
        if suffix not in SUPPORTED_EDIT_EXTENSIONS:
            logger.warning("read_for_edit: unsupported type %s for %s", suffix, file_path)
            return None
        try:
            if suffix == ".txt":
                content = path.read_text(encoding="utf-8", errors="replace")
            else:
                content = self._read_docx(str(path))
            if len(content) > MAX_READ_CHARS:
                content = content[:MAX_READ_CHARS] + "\n\n[... document truncated for length ...]"
            return content
        except Exception as exc:
            logger.error("read_for_edit failed for %s: %s", file_path, exc)
            return None

    def _read_docx(self, file_path: str) -> str:
        from docx import Document
        doc = Document(file_path)
        parts: List[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text.strip()
                        if text:
                            parts.append(text)
        return "\n\n".join(parts)

    # ------------------------------------------------------------------
    # Generate proposal
    # ------------------------------------------------------------------

    async def generate_proposal(
        self,
        file_path: str,
        content: str,
        instruction: str,
        llm_service: Any,
    ) -> Optional[EditProposal]:
        """Ask the LLM to produce a structured find/replace diff, then validate it."""
        prompt = (
            "You are a precise document editor. Given a document and an edit instruction, "
            "output a JSON array of find/replace changes.\n\n"
            "Rules:\n"
            '  - Each "find" MUST be an EXACT verbatim substring of the document '
            "(copy character-for-character, including punctuation and spacing)\n"
            "  - Make only the minimum changes to fulfill the instruction\n"
            '  - To REMOVE text: set "replace" to an empty string ""\n'
            "  - If you cannot find the exact text, output an empty array []\n"
            "  - Output ONLY the JSON array — no explanation, no markdown fences\n\n"
            f"Document:\n---\n{content}\n---\n\n"
            f"Instruction: {instruction}\n\n"
            "JSON array examples:\n"
            '  Replace: [{"find": "old text", "replace": "new text"}]\n'
            '  Remove:  [{"find": "text to delete", "replace": ""}]\n\n'
            "JSON array:\n"
        )
        try:
            raw = await llm_service.generate_simple(
                prompt,
                prompt_type="short_direct",
                max_completion_tokens=512,
            )
        except Exception as exc:
            logger.error("generate_proposal LLM call failed: %s", exc)
            return None

        changes = self._parse_and_validate_changes(raw, content)
        if not changes:
            logger.info("generate_proposal: no valid changes produced for %s", file_path)
            return None

        proposal = EditProposal(
            proposal_id=str(uuid.uuid4()),
            document_name=Path(file_path).name,
            file_path=file_path,
            file_type=Path(file_path).suffix.lower().lstrip("."),
            changes=changes,
            summary=self._make_summary(changes),
        )
        self.store_proposal(proposal)
        return proposal

    def _parse_and_validate_changes(
        self, raw: str, content: str
    ) -> List[EditChange]:
        match = re.search(r"\[.*\]", raw, re.DOTALL)
        if not match:
            logger.warning("edit proposal: no JSON array in LLM output")
            return []
        try:
            items = json.loads(match.group(0))
        except json.JSONDecodeError as exc:
            logger.warning("edit proposal: JSON parse error: %s", exc)
            return []
        changes: List[EditChange] = []
        for item in items:
            if not isinstance(item, dict):
                continue
            find = str(item.get("find", "")).strip()
            replace = str(item.get("replace") or "")
            if not find:
                continue
            if find not in content:
                logger.warning(
                    "edit proposal: 'find' text absent from document, skipping: %r",
                    find[:80],
                )
                continue
            changes.append(EditChange(find=find, replace=replace))
        return changes

    @staticmethod
    def _make_summary(changes: List[EditChange]) -> str:
        n = len(changes)
        if n == 0:
            return "No changes proposed"
        if n == 1:
            preview = changes[0].find[:50].replace("\n", " ")
            ellipsis = "…" if len(changes[0].find) > 50 else ""
            return f'1 change: "{preview}{ellipsis}"'
        return f"{n} changes proposed"

    # ------------------------------------------------------------------
    # Apply
    # ------------------------------------------------------------------

    def apply_proposal(self, proposal: EditProposal) -> Dict[str, Any]:
        """
        Validate, back up, write the file.

        Returns dict with keys: applied_changes (int), backup_path (str).
        Raises ValueError on safety / IO errors — caller should catch and 500.
        """
        current_dir = self.current_directory
        if not current_dir:
            raise ValueError("No directory is currently loaded")

        file_path = proposal.file_path

        # Scope guard: file must be under the indexed directory
        try:
            abs_file = os.path.abspath(file_path)
            abs_dir = os.path.abspath(current_dir)
            if not abs_file.startswith(abs_dir + os.sep) and abs_file != abs_dir:
                raise ValueError("File is outside the indexed directory")
        except Exception as exc:
            raise ValueError(f"Path validation failed: {exc}")

        path = Path(file_path)
        if not path.exists() or not path.is_file():
            raise ValueError(f"File no longer exists: {file_path}")

        suffix = path.suffix.lower()
        if suffix not in SUPPORTED_EDIT_EXTENSIONS:
            raise ValueError(f"Unsupported file type for editing: {suffix}")

        # Timestamped backup (same directory, invisible to the user unless they look)
        ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        backup_path = path.with_suffix(f".{ts}.bak{suffix}")
        shutil.copy2(str(path), str(backup_path))
        logger.info("Backup created: %s", backup_path)

        try:
            if suffix == ".txt":
                applied = self._apply_txt(str(path), proposal.changes)
            else:
                applied = self._apply_docx(str(path), proposal.changes)
        except Exception as exc:
            # Restore from backup on failure
            shutil.copy2(str(backup_path), str(path))
            logger.error("Apply failed, restored from backup: %s", exc)
            raise ValueError(f"Apply failed — file restored: {exc}")

        return {"applied_changes": applied, "backup_path": str(backup_path)}

    def _apply_txt(self, file_path: str, changes: List[EditChange]) -> int:
        content = Path(file_path).read_text(encoding="utf-8", errors="replace")
        applied = 0
        for change in changes:
            count = content.count(change.find)
            if count:
                content = content.replace(change.find, change.replace)
                applied += count
            else:
                logger.warning(
                    "_apply_txt: find text no longer present, skipping: %r", change.find[:60]
                )
        Path(file_path).write_text(content, encoding="utf-8")
        return applied

    def _apply_docx(self, file_path: str, changes: List[EditChange]) -> int:
        """
        Apply changes to a DOCX file.

        For each paragraph (and table cell) concatenate all run texts, apply ALL
        replacements, then write the result back to the first run and blank the
        remaining runs.  Preserves paragraph-level formatting; collapses per-run
        formatting inside changed paragraphs — acceptable for Phase 1.
        """
        from docx import Document

        doc = Document(file_path)
        applied = 0

        def _try_replace_in_para(para: Any) -> None:
            nonlocal applied
            para_text = para.text
            new_text = para_text
            for change in changes:
                count = new_text.count(change.find)
                if count:
                    new_text = new_text.replace(change.find, change.replace)
                    applied += count
            if new_text != para_text and para.runs:
                para.runs[0].text = new_text
                for run in para.runs[1:]:
                    run.text = ""

        for para in doc.paragraphs:
            _try_replace_in_para(para)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _try_replace_in_para(para)

        doc.save(file_path)
        return applied

    # ------------------------------------------------------------------
    # Direct save (TipTap editor content → disk)
    # ------------------------------------------------------------------

    def save_content(self, file_path: str, content: str, fmt: str) -> Dict[str, Any]:
        """
        Write full document content coming from the TipTap editor.

        fmt must be "txt" or "docx".
        Backs up the original file before overwriting.
        Returns dict with backup_path.
        Raises ValueError on validation / IO errors.
        """
        current_dir = self.current_directory
        if not current_dir:
            raise ValueError("No directory is currently loaded")

        try:
            abs_file = os.path.abspath(file_path)
            abs_dir = os.path.abspath(current_dir)
            if not abs_file.startswith(abs_dir + os.sep) and abs_file != abs_dir:
                raise ValueError("File is outside the indexed directory")
        except Exception as exc:
            raise ValueError(f"Path validation failed: {exc}")

        path = Path(file_path)
        if not path.exists() or not path.is_file():
            raise ValueError(f"File no longer exists: {file_path}")

        suffix = path.suffix.lower()
        if suffix not in SUPPORTED_EDIT_EXTENSIONS:
            raise ValueError(f"Unsupported file type for editing: {suffix}")

        ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        backup_path = path.with_suffix(f".{ts}.bak{suffix}")
        shutil.copy2(str(path), str(backup_path))
        logger.info("Backup created: %s", backup_path)

        if fmt != "txt":
            raise ValueError("Only .txt files support direct content save. Use AI proposals to edit DOCX.")

    # ------------------------------------------------------------------
    # Cell-level save (Excel)
    # ------------------------------------------------------------------

    def save_cells(self, file_path: str, changes: List[Dict[str, Any]]) -> Dict[str, Any]:
        """
        Apply cell-level edits to an XLSX file.

        changes: list of {sheet: str, cell: str (e.g. "B5"), value: str}
        Only .xlsx is supported — .xls is read-only via xlrd.
        Backs up the file before writing.
        """
        current_dir = self.current_directory
        if not current_dir:
            raise ValueError("No directory is currently loaded")

        try:
            abs_file = os.path.abspath(file_path)
            abs_dir = os.path.abspath(current_dir)
            if not abs_file.startswith(abs_dir + os.sep) and abs_file != abs_dir:
                raise ValueError("File is outside the indexed directory")
        except Exception as exc:
            raise ValueError(f"Path validation failed: {exc}")

        path = Path(file_path)
        if not path.exists() or not path.is_file():
            raise ValueError(f"File no longer exists: {file_path}")

        if path.suffix.lower() != ".xlsx":
            raise ValueError(
                "Cell editing is only supported for .xlsx files. "
                ".xls files are read-only — convert to .xlsx to edit."
            )

        ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        backup_path = path.with_suffix(f".{ts}.bak.xlsx")
        shutil.copy2(str(path), str(backup_path))
        logger.info("Backup created: %s", backup_path)

        try:
            from openpyxl import load_workbook
            wb = load_workbook(str(path))
            applied = 0
            for change in changes:
                sheet_name = change.get("sheet", "")
                cell_address = change.get("cell", "")
                raw_value = change.get("value", "")
                if not sheet_name or not cell_address:
                    continue
                if sheet_name not in wb.sheetnames:
                    logger.warning("save_cells: sheet %r not found, skipping", sheet_name)
                    continue
                wb[sheet_name][cell_address] = self._coerce_cell_value(raw_value)
                applied += 1
            wb.save(str(path))
            wb.close()
        except Exception as exc:
            shutil.copy2(str(backup_path), str(path))
            logger.error("save_cells failed, restored from backup: %s", exc)
            raise ValueError(f"Cell save failed — file restored: {exc}")

        return {"applied_changes": applied, "backup_path": str(backup_path)}

    @staticmethod
    def _coerce_cell_value(raw: str):
        """Parse as int or float when possible, else return as string."""
        stripped = raw.strip()
        if not stripped:
            return None
        try:
            return int(stripped)
        except ValueError:
            pass
        try:
            return float(stripped)
        except ValueError:
            pass
        return stripped

        try:
            path.write_text(content, encoding="utf-8")
        except Exception as exc:
            shutil.copy2(str(backup_path), str(path))
            logger.error("save_content failed, restored from backup: %s", exc)
            raise ValueError(f"Save failed — file restored: {exc}")

        return {"backup_path": str(backup_path)}

    def _save_html_as_docx(self, file_path: str, html: str) -> None:
        """Convert TipTap HTML to a DOCX file using python-docx."""
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Remove default empty paragraph Word adds
        for p in doc.paragraphs:
            p._element.getparent().remove(p._element)

        parser = _TipTapHTMLParser(doc)
        parser.feed(html)
        doc.save(file_path)


class _TipTapHTMLParser(HTMLParser):
    """
    Minimal HTML → python-docx converter for TipTap StarterKit output.
    Handles: p, h1-h6, ul/ol/li, strong, em, u, s, br, blockquote, pre/code.
    """

    _HEADING_LEVELS = {"h1": 1, "h2": 2, "h3": 3, "h4": 4, "h5": 5, "h6": 6}

    def __init__(self, doc: Any) -> None:
        super().__init__()
        self._doc = doc
        self._para: Any = None          # current paragraph
        self._run_stack: List[Dict] = []  # formatting context stack
        self._list_stack: List[str] = []  # "ul" | "ol"
        self._list_counters: List[int] = []
        self._in_pre = False

    # ------------------------------------------------------------------ helpers

    def _current_fmt(self) -> Dict:
        fmt: Dict = {}
        for layer in self._run_stack:
            fmt.update(layer)
        return fmt

    def _new_para(self, style: str = "Normal") -> Any:
        from docx.oxml.ns import qn
        p = self._doc.add_paragraph(style=style)
        self._para = p
        return p

    def _add_run(self, text: str) -> None:
        if self._para is None:
            self._new_para()
        fmt = self._current_fmt()
        run = self._para.add_run(text)
        run.bold = fmt.get("bold", False)
        run.italic = fmt.get("italic", False)
        run.underline = fmt.get("underline", False)
        if fmt.get("strike"):
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            rPr = run._r.get_or_add_rPr()
            strike = OxmlElement("w:strike")
            rPr.append(strike)

    # ------------------------------------------------------------------ tags

    def handle_starttag(self, tag: str, attrs: list) -> None:
        tag = tag.lower()

        if tag in self._HEADING_LEVELS:
            level = self._HEADING_LEVELS[tag]
            self._new_para(style=f"Heading {level}")

        elif tag == "p":
            if self._list_stack:
                # Inside a list item — paragraph already created by <li>
                pass
            else:
                self._new_para()

        elif tag == "li":
            list_type = self._list_stack[-1] if self._list_stack else "ul"
            if list_type == "ol":
                self._list_counters[-1] += 1
                self._new_para(style="List Number")
            else:
                self._new_para(style="List Bullet")

        elif tag == "ul":
            self._list_stack.append("ul")
            self._list_counters.append(0)

        elif tag == "ol":
            self._list_stack.append("ol")
            self._list_counters.append(0)

        elif tag == "blockquote":
            self._new_para(style="Quote")

        elif tag in ("pre", "code"):
            self._in_pre = True
            if tag == "pre":
                self._new_para(style="Normal")

        elif tag == "br":
            if self._para is not None:
                self._para.add_run().add_break()

        elif tag == "strong" or tag == "b":
            self._run_stack.append({"bold": True})

        elif tag == "em" or tag == "i":
            self._run_stack.append({"italic": True})

        elif tag == "u":
            self._run_stack.append({"underline": True})

        elif tag in ("s", "del", "strike"):
            self._run_stack.append({"strike": True})

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()

        if tag in self._HEADING_LEVELS or tag == "p":
            self._para = None

        elif tag == "li":
            self._para = None

        elif tag in ("ul", "ol"):
            if self._list_stack:
                self._list_stack.pop()
                self._list_counters.pop()

        elif tag in ("pre", "code"):
            self._in_pre = False
            if tag == "pre":
                self._para = None

        elif tag in ("strong", "b", "em", "i", "u", "s", "del", "strike"):
            if self._run_stack:
                self._run_stack.pop()

    def handle_data(self, data: str) -> None:
        if not data:
            return
        if self._para is None and not self._list_stack:
            self._new_para()
        if self._para is not None:
            self._add_run(data)


